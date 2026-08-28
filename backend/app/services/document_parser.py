import re
from pathlib import Path
from typing import List, Dict, Any
import pypdf
import docx

def clean_text(text: str) -> str:
    """Normalize whitespace and remove non-printable characters while preserving punctuation and numbers."""
    if not text:
        return ""
    # Replace multiple newlines/tabs with standardized spaces
    text = re.sub(r'[\r\f\v]', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()

def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Extract page-by-page text from a PDF file."""
    pages_data = []
    reader = pypdf.PdfReader(file_path)
    for index, page in enumerate(reader.pages, start=1):
        extracted = page.extract_text() or ""
        cleaned = clean_text(extracted)
        if cleaned:
            pages_data.append({
                "page_number": index,
                "text": cleaned
            })
    return pages_data

def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from a DOCX file, grouping by sections/paragraphs."""
    doc = docx.Document(file_path)
    full_paragraphs = []
    
    # Extract body paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            full_paragraphs.append(p.text.strip())
            
    # Extract tables content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_paragraphs.append(row_text)

    # In DOCX without explicit pagination, segment approximately every 400 words as a virtual page
    pages_data = []
    current_page = 1
    current_text_buf = []
    current_word_count = 0

    for para in full_paragraphs:
        words = len(para.split())
        current_text_buf.append(para)
        current_word_count += words
        
        if current_word_count >= 350:
            pages_data.append({
                "page_number": current_page,
                "text": clean_text("\n\n".join(current_text_buf))
            })
            current_page += 1
            current_text_buf = []
            current_word_count = 0

    if current_text_buf:
        pages_data.append({
            "page_number": current_page,
            "text": clean_text("\n\n".join(current_text_buf))
        })

    return pages_data

def parse_txt(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from a TXT file, segmenting into logical pages by length or section markers."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    cleaned = clean_text(content)
    # Split into sections if page markers like '--- Page X ---' or '=== Page X ===' exist
    page_markers = list(re.finditer(r'(?:---+|\*\*\*+|===+)\s*(?:Page|PAGE|Section)\s*(\d+)?\s*(?:---+|\*\*\*+|===+)', cleaned))
    
    if page_markers:
        pages_data = []
        last_pos = 0
        for i, match in enumerate(page_markers):
            start = match.start()
            if start > last_pos:
                seg_text = cleaned[last_pos:start].strip()
                if seg_text:
                    pages_data.append({
                        "page_number": len(pages_data) + 1,
                        "text": seg_text
                    })
            last_pos = match.end()
        if last_pos < len(cleaned):
            seg_text = cleaned[last_pos:].strip()
            if seg_text:
                pages_data.append({
                    "page_number": len(pages_data) + 1,
                    "text": seg_text
                })
        return pages_data

    # Otherwise segment into virtual pages (~400 words each)
    paragraphs = cleaned.split("\n\n")
    pages_data = []
    current_page = 1
    current_buf = []
    current_words = 0

    for p in paragraphs:
        if not p.strip():
            continue
        words = len(p.split())
        current_buf.append(p.strip())
        current_words += words
        if current_words >= 350:
            pages_data.append({
                "page_number": current_page,
                "text": "\n\n".join(current_buf)
            })
            current_page += 1
            current_buf = []
            current_words = 0

    if current_buf:
        pages_data.append({
            "page_number": current_page,
            "text": "\n\n".join(current_buf)
        })

    return pages_data if pages_data else [{"page_number": 1, "text": cleaned}]

def extract_document_pages(file_path: str, file_type: str) -> List[Dict[str, Any]]:
    """Universal parser dispatching according to file extension."""
    ft = file_type.lower().replace(".", "")
    if ft == "pdf":
        return parse_pdf(file_path)
    elif ft in ["docx", "doc"]:
        return parse_docx(file_path)
    elif ft in ["txt", "md", "csv"]:
        return parse_txt(file_path)
    else:
        # Default text fallback
        return parse_txt(file_path)
